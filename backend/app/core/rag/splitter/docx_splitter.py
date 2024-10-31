from docx import Document
from app.core.ocr.ocr_model import OCR_Model
from langchain_core.documents import Document as LangChainDocument
from app.core.rag.splitter.structured_file import TextSplitter
from config.config import SPPLITTER_MODEL
from config.splitter_model import SplitterModel
from typing import List



class DocxSplitter(TextSplitter):
    def __init__(self, file_path: str,splitter_args=None, splitter_model: SplitterModel = SPPLITTER_MODEL, *args, **kwargs):
        super().__init__(splitter_args=splitter_args,SPPLITTER_MODEL=splitter_model,*args, **kwargs)
        self.ocr_model = OCR_Model()
        self.file_path = file_path
        # self.splitter_pattern = splitter_pattern
        self.splitter_model = splitter_model

    def load(self):
        # 读取 docx 文件
        doc = Document(self.file_path)
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text + "\n"
            # print(para.text)
        for table in doc.tables:
            # 获取表格行数与列数
            num_cols = len(table.columns)
            # 创建表头分隔行
            full_text += "|" + " | ".join(["---"] * num_cols) + "|\n"
            
            for row in table.rows:
                # 逐行读取单元格内容并添加分隔符
                row_text = "| " + " | ".join(cell.text.strip() for cell in row.cells) + " |\n"
                full_text += row_text
            full_text += "\n\n"

        print(full_text)
        
        return full_text
    def split(self)->List[LangChainDocument]:
        full_text = self.load()
        return super().split(full_text)

if __name__ == "__main__":
    docx_splitter = DocxSplitter("/Users/markyangkp/Documents/信息整理/常用校园信息集合.docx")
    docx_splitter.load()